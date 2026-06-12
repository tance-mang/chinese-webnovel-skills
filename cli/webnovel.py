#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
网文工坊 · 终端版
独立于 Claude 运行，接任意 OpenAI 兼容 API（DeepSeek / Kimi / 硅基流动 / OpenRouter / 本地 Ollama / OpenAI ...）。

用法：
  python webnovel.py                      # 进入对话模式
  python webnovel.py "用 hook 写个修真开头"   # 一次性
  python webnovel.py --full               # 加载完整技能+知识库(需大上下文模型)
  python webnovel.py --list               # 列出所有技能

对话模式里输入 /help 看全部命令（存稿/换格式/学习模式/校对过审/历史回溯）。

配置（三选一）：
  1) 环境变量：WN_API_KEY / WN_BASE_URL / WN_MODEL
  2) cli/config.json（复制 config.example.json 改）
  3) 启动后按提示填
存稿偏好（书名/格式/目录/自动记录）写在 config.json，可在对话里用 /set 改，不写死。
依赖：pip install openai   （存 Word 另需：pip install python-docx）
"""
import os
import sys
import json
import argparse
import datetime

HERE = os.path.dirname(os.path.abspath(__file__))
EXPORTS = os.path.join(HERE, "..", "exports")
CONFIG_PATH = os.path.join(HERE, "config.json")

# 仅这些是“存稿偏好”，会被 /set 持久化进 config.json（绝不写 api_key）
PREF_KEYS = ("book", "output_dir", "save_format", "auto_log")
SAVE_FORMATS = ("txt", "md", "docx")


def read(path):
    with open(path, encoding="utf-8") as f:
        return f.read()


def load_config():
    cfg = {}
    if os.path.exists(CONFIG_PATH):
        try:
            cfg = json.load(open(CONFIG_PATH, encoding="utf-8"))
        except Exception as e:
            print("[警告] config.json 解析失败：", e)
    # API（环境变量优先）
    cfg["base_url"] = os.environ.get("WN_BASE_URL", cfg.get("base_url") or "https://api.deepseek.com/v1")
    cfg["api_key"] = os.environ.get("WN_API_KEY", cfg.get("api_key") or "")
    cfg["model"] = os.environ.get("WN_MODEL", cfg.get("model") or "deepseek-chat")
    # 存稿偏好（给默认值，不写死——都能在对话里 /set 改）
    cfg["book"] = cfg.get("book") or "未命名"
    cfg["output_dir"] = cfg.get("output_dir") or os.path.join(os.getcwd(), "output")
    fmt = (cfg.get("save_format") or "txt").lower()
    cfg["save_format"] = fmt if fmt in SAVE_FORMATS else "txt"
    cfg["auto_log"] = cfg.get("auto_log", True)
    if not cfg["api_key"]:
        print("\n还没配 API Key。终端版是个壳，需要接一个 AI 才能写作")
        print("（DeepSeek / Kimi / 硅基流动 / OpenRouter 等任意一家，有免费额度，去哪弄见 cli/README.md「第四步」）。")
        print("拿到后两种填法：① 把 config.example.json 复制成 config.json 填进去（一劳永逸）；② 现在直接粘贴下面：")
        cfg["api_key"] = input("  粘贴 API Key（还没有就直接回车，等接好 AI 再来）：").strip()
        if not cfg["api_key"]:
            print("（已跳过。配好 key 再运行 python webnovel.py 即可。）")
            sys.exit(0)
    return cfg


def save_prefs(cfg):
    """把存稿偏好写回 config.json，保留其它字段（含 api_key 原样，不新增）。"""
    data = {}
    if os.path.exists(CONFIG_PATH):
        try:
            data = json.load(open(CONFIG_PATH, encoding="utf-8"))
        except Exception:
            data = {}
    for k in PREF_KEYS:
        data[k] = cfg.get(k)
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print("[警告] 偏好保存失败：", e)
        return False


def build_system_prompt(full=False):
    parts = [read(os.path.join(EXPORTS, "instructions.md"))]
    if full:
        parts.append("\n\n# 技能工作流（详细）\n" + read(os.path.join(EXPORTS, "knowledge-skills.md")))
        parts.append("\n\n# 知识库（详细）\n" + read(os.path.join(EXPORTS, "knowledge-references.md")))
    else:
        parts.append("\n\n> 提示：如需某技能的完整工作流或知识库细节，按你已知的网文写作方法执行；"
                     "用 --full 启动可加载全部技能 + 知识库的完整内容。")
    return "\n".join(parts)


COACH_DEPTHS = {
    "1": "用【速查清单】式：每次产出后只列 3-5 个要点（钩子/爽点/节奏/人设），最快。",
    "2": "用【详细讲解】：每次产出后讲清‘为什么这么写’的原理 + 1 个例子，通俗不堆术语。",
    "3": "用【边写边教】：先正常写，再回过头拆解这段为什么这么处理、哪里还能更好。",
}


def coach_addendum(depth):
    return ("\n\n【学习模式已开启】你现在兼任网文写作教练，面向新手：" + COACH_DEPTHS.get(depth, COACH_DEPTHS["2"]) +
            " 指出可改进处并鼓励作者动手。作者输入 /coach off 即回到正常写作模式。")


def list_skills():
    txt = read(os.path.join(EXPORTS, "knowledge-skills.md"))
    print("可用技能（在对话里直接说需求或点名技能）：")
    for line in txt.splitlines():
        if line.startswith("## 技能："):
            print("  -", line.replace("## 技能：", "").strip())


def chat_stream(client, model, messages):
    """流式输出，返回完整回复。"""
    full = []
    resp = client.chat.completions.create(model=model, messages=messages, stream=True, temperature=0.8)
    for chunk in resp:
        delta = chunk.choices[0].delta.content if chunk.choices and chunk.choices[0].delta else None
        if delta:
            print(delta, end="", flush=True)
            full.append(delta)
    print()
    return "".join(full)


# ---------- 存稿 / 回溯 ----------

def sanitize(name):
    bad = '\\/:*?"<>|'
    return "".join("_" if c in bad else c for c in (name or "")).strip() or "未命名"


def book_dir(cfg):
    return os.path.join(cfg["output_dir"], sanitize(cfg["book"]))


def unique_path(path):
    """文件已存在就自增 _2/_3…，绝不覆盖（不丢稿）。"""
    if not os.path.exists(path):
        return path
    root, ext = os.path.splitext(path)
    i = 2
    while os.path.exists(f"{root}_{i}{ext}"):
        i += 1
    return f"{root}_{i}{ext}"


def write_text_file(path, title, text, fmt):
    if fmt == "docx":
        try:
            from docx import Document
        except ImportError:
            print("[提示] 未装 python-docx，本次改存 .txt（装 Word 支持：pip install python-docx）")
            return write_text_file(os.path.splitext(path)[0] + ".txt", title, text, "txt")
        doc = Document()
        if title:
            doc.add_heading(title, level=1)
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(path)
    else:
        body = (("# " + title + "\n\n") if (fmt == "md" and title) else
                ((title + "\n\n") if title else "")) + text
        with open(path, "w", encoding="utf-8") as f:
            f.write(body)
    return path


def save_chapter(cfg, text, name=None):
    """把正文存进 书文件夹/正文/，按当前格式，命名不覆盖，返回绝对路径。"""
    name = sanitize(name) if name else "第" + datetime.datetime.now().strftime("%m%d_%H%M") + "章"
    folder = os.path.join(book_dir(cfg), "正文")
    os.makedirs(folder, exist_ok=True)
    path = unique_path(os.path.join(folder, f"{name}.{cfg['save_format']}"))
    write_text_file(path, name, text, cfg["save_format"])
    log_history(cfg, "save", path, len(text))
    return os.path.abspath(path)


def append_session_log(cfg, role, text):
    """自动记录：每轮都进 创作记录/对话记录.txt，永不丢。"""
    if not cfg.get("auto_log"):
        return
    folder = os.path.join(book_dir(cfg), "创作记录")
    os.makedirs(folder, exist_ok=True)
    p = os.path.join(folder, "对话记录.txt")
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(p, "a", encoding="utf-8") as f:
        f.write(f"\n----- {role} · {ts} -----\n{text}\n")


def log_history(cfg, action, path, wordcount):
    """创作记录索引（类 git 的可回溯清单）：时间/动作/文件/字数。"""
    folder = os.path.join(book_dir(cfg), "创作记录")
    os.makedirs(folder, exist_ok=True)
    p = os.path.join(folder, "history.tsv")
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(p, "a", encoding="utf-8") as f:
        f.write(f"{ts}\t{action}\t{os.path.abspath(path)}\t{wordcount}\n")


def show_history(cfg):
    p = os.path.join(book_dir(cfg), "创作记录", "history.tsv")
    if not os.path.exists(p):
        print("(还没有存稿记录。用 /save 存一章试试)")
        return
    print(f"《{cfg['book']}》创作记录（可回溯，文件都在：{os.path.abspath(book_dir(cfg))}）：")
    for line in read(p).splitlines():
        cols = line.split("\t")
        if len(cols) == 4:
            ts, act, path, wc = cols
            print(f"  {ts} | {act} | {wc}字 | {os.path.basename(path)}")


def _ask(client, cfg, messages, state, prompt):
    """发一条 user 消息、流式取回复、记进历史与 last_reply。"""
    messages.append({"role": "user", "content": prompt})
    print("工坊 > ", end="", flush=True)
    try:
        reply = chat_stream(client, cfg["model"], messages)
        messages.append({"role": "assistant", "content": reply})
        state["last_reply"] = reply
        append_session_log(cfg, "工坊", reply)
        return reply
    except Exception as e:
        print("\n[错误]", e)
        messages.pop()
        return None


COWRITE_FMT = ("给：推荐版本★(评分0-100 + 推荐原因 + 风险)，再给 备选A、备选B(各带评分+一句特点)。"
               "正文和分析分区，每版用 ━━━ 分隔，正文放【正文】下。守网文手感：短句、show不tell、"
               "对话用弯双引号\"\"不用「」、破折号极少。末尾列：1用推荐 / 2用A / 3用B / r重写。")

# 模式化入口：把 30+ 技能收进 4 个门，背后自动调技能（LLM 已加载全部技能工作流）
MODE_MENU = """你现在在哪一步？(输下面命令进入；也可直接说人话)
  /new        新建小说  选题→(强设定搭世界)→金手指→人设→大纲→开篇钩→第一章
  /continue   继续写    读档回顾→写下一章/共写→查一致→更新档案
  /optimize   优化修改  追读体检→去AI味/打分→情绪演出来→错别字敏感词校对
  /submit     投稿分析  字数门槛→三章签约体检→过审校对→平台趋势→书名简介"""

MODE_KICKOFF = {
    "new": ("我要【新建一本小说】。按入口流程带我走：先问我 频道/题材/平台/篇幅，再一步步 "
            "选题→(强设定就先搭世界)→金手指→人设→大纲→开篇钩→第一章。一次只推进一步，每步问我要的信息，别一股脑全做完。"),
    "continue": ("我要【继续写之前的书】。先帮我回顾'写到哪了/上章结尾钩/待回收伏笔/主角与情绪状态'"
                 "(我会贴档案或正文)，再写下一章(可用共写)；写完提醒我更新档案、查前后一致。一次一步。"),
    "optimize": ("我要【优化/修改已有正文】。我把正文发你后，按 追读体检→去AI味/量化打分→情绪演出来→"
                 "错别字敏感词校对 的顺序过，一项一项来，先做最影响读者留存的。"),
    "submit": ("我要【投稿/分析】。帮我核对 篇幅字数门槛→三章签约体检→过审校对→平台趋势→书名简介标签。"
               "先问我目标平台和现在的字数。"),
}


def cowrite_pick(cfg, client, messages, state, depth=0):
    """多版本生成后，让作者选 1/2/3 采用并存档，r 重写。"""
    if depth > 3:
        return
    try:
        choice = input("选 1/2/3 采用并存档 ｜ r 重写 ｜ 回车跳过：").strip().lower()
    except (EOFError, KeyboardInterrupt):
        return
    if choice in ("1", "2", "3"):
        label = {"1": "推荐版本", "2": "备选A", "3": "备选B"}[choice]
        reply = _ask(client, cfg, messages, state,
                     f"只输出刚才那个【{label}】的正文本身，不要标题/评分/分析/分隔线/任何其它字。")
        if reply:
            path = save_chapter(cfg, reply, None)
            log_history(cfg, "共写采用", path, len(reply))
            print(f"✅ 已采用并存：{path}")
    elif choice == "r":
        _ask(client, cfg, messages, state, "这几版不满意，换个思路重写这几个版本（同样给推荐+备选两个）。")
        cowrite_pick(cfg, client, messages, state, depth + 1)


HELP = """命令一览（直接输文字=继续写；命令以 / 开头）：
  /help              本帮助
  /mode              四个入口菜单（不知道用啥先看这个）
  /new /continue /optimize /submit   新建 / 继续写 / 优化 / 投稿（自动走流程调技能）
  /co [我写的一段]    共写·辅助：续写下一段，给 推荐+备选 选数字采用（卡在“下一句写啥”用这个）
  /director          共写·导演：填 目标/冲突/情绪/节奏，AI 出多版正文
  /write [需求]       自动写：按章纲/需求直接出整段正文
  /save [章名]        把【上一条回复】存成正文文件（不覆盖、存完显示路径）
  /book <书名>        设置当前书名（决定存到哪个书文件夹）
  /set format txt|md|docx   选存稿格式（默认 txt）
  /set dir <路径>     选存到哪个目录
  /set autolog on|off 自动记录每轮对话(永不丢) 开/关
  /config            看当前设置（书名/目录/格式/模型）
  /history           看创作记录、回溯写过什么（类 git）
  /proofread         校对上一条回复：错别字+敏感词，给 原词→建议词
  /coach [off]       学习模式 开/关
  /list              列出全部技能
  exit / quit / 退出  离开
"""


def print_config(cfg):
    print(f"当前设置：书名《{cfg['book']}》｜格式 .{cfg['save_format']}｜自动记录 {'开' if cfg['auto_log'] else '关'}")
    print(f"存稿目录：{os.path.abspath(book_dir(cfg))}")
    print(f"模型：{cfg['model']} @ {cfg['base_url']}")


def handle_command(line, cfg, client, messages, state):
    """返回 True 表示已处理掉这条命令（不发给模型）。"""
    parts = line.split()
    cmd = parts[0].lower()
    arg = line[len(parts[0]):].strip()

    if cmd == "/help":
        print(HELP)
    elif cmd == "/config":
        print_config(cfg)
    elif cmd == "/list":
        list_skills()
    elif cmd == "/book":
        if not arg:
            print("用法：/book 书名")
        else:
            cfg["book"] = arg
            save_prefs(cfg)
            print(f"当前书名设为《{cfg['book']}》，存稿目录：{os.path.abspath(book_dir(cfg))}")
    elif cmd == "/set":
        sp = arg.split(None, 1)
        key = sp[0].lower() if sp else ""
        val = sp[1].strip() if len(sp) > 1 else ""
        if key == "format" and val.lower() in SAVE_FORMATS:
            cfg["save_format"] = val.lower(); save_prefs(cfg); print(f"存稿格式设为 .{cfg['save_format']}")
        elif key == "dir" and val:
            cfg["output_dir"] = val; save_prefs(cfg); print(f"存稿目录设为：{os.path.abspath(book_dir(cfg))}")
        elif key == "autolog" and val.lower() in ("on", "off", "开", "关"):
            cfg["auto_log"] = val.lower() in ("on", "开"); save_prefs(cfg)
            print(f"自动记录已{'开启' if cfg['auto_log'] else '关闭'}")
        else:
            print("用法：/set format txt|md|docx ｜ /set dir <路径> ｜ /set autolog on|off")
    elif cmd == "/save":
        if not state.get("last_reply"):
            print("(还没有可保存的回复。先让我写点什么)")
        else:
            path = save_chapter(cfg, state["last_reply"], arg or None)
            print(f"✅ 已存：{path}")
    elif cmd == "/history":
        show_history(cfg)
    elif cmd == "/mode":
        print(MODE_MENU)
    elif cmd in ("/new", "/continue", "/optimize", "/submit"):
        _ask(client, cfg, messages, state, MODE_KICKOFF[cmd[1:]])
    elif cmd == "/co":
        ctx = arg or state.get("last_reply", "")
        base = ("【共写·辅助模式】接着前文续写下一段（约 200–400 字）。" + COWRITE_FMT +
                "\n\n前文 / 我刚写的：\n" + (ctx if ctx else "(还没正文，请基于本书设定开个头)"))
        if _ask(client, cfg, messages, state, base):
            cowrite_pick(cfg, client, messages, state)
    elif cmd == "/director":
        print("导演模式：你定方向，AI 出多版正文。回车跳过某项。")
        try:
            goal = input("  目标（这段要达成什么）：").strip()
            conflict = input("  冲突 / 阻碍：").strip()
            emotion = input("  情绪基调：").strip()
            pace = input("  节奏（快 / 缓）：").strip()
        except (EOFError, KeyboardInterrupt):
            print(); return True
        spec = f"目标:{goal or '不限'}；冲突:{conflict or '不限'}；情绪:{emotion or '不限'}；节奏:{pace or '不限'}"
        base = ("【共写·导演模式】作者只给方向，你写正文。按下面方向写本段（约 300–500 字），" +
                COWRITE_FMT + "\n\n方向：" + spec)
        if _ask(client, cfg, messages, state, base):
            cowrite_pick(cfg, client, messages, state)
    elif cmd == "/write":
        base = ("【自动写】按本书章纲 / 进度直接写下一段可发布正文。" +
                ("需求：" + arg if arg else "没给具体需求就按大纲推进下一章。") +
                " 守网文手感：短句、show不tell、对话用弯双引号\"\"不用「」、章尾留钩。")
        _ask(client, cfg, messages, state, base)
        print("(满意就 /save 存档)")
    elif cmd == "/proofread":
        if not state.get("last_reply"):
            print("(没有可校对的内容。先写一段)")
        else:
            q = ("请用 proofread 技能校对下面这段网文：①错别字(按句意，给 原词→建议词) "
                 "②可能触发平台审核的敏感/违禁内容(标类别+替换建议，说明是启发式非平台确切词库)。"
                 "逐条列‘原句→建议’：\n\n" + state["last_reply"])
            messages.append({"role": "user", "content": q})
            print("工坊 > ", end="", flush=True)
            try:
                reply = chat_stream(client, cfg["model"], messages)
                messages.append({"role": "assistant", "content": reply})
                append_session_log(cfg, "校对", reply)
            except Exception as e:
                print("\n[错误]", e); messages.pop()
    elif cmd == "/coach":
        if arg.lower() == "off":
            if state.get("coach"):
                messages[0]["content"] = state["base_system"]
                state["coach"] = False
                print("已退出学习模式，回到正常写作。")
            else:
                print("当前不在学习模式。")
        else:
            print("进入学习模式。能教你：开篇钩子/爽点打脸/节奏呼吸/人设弧光/去AI味/投稿过稿 等。")
            print("怎么学？ 1) 速查清单(最快)  2) 详细讲解(原理+例子)  3) 边写边教")
            try:
                depth = input("选 1/2/3（回车默认2）：").strip() or "2"
            except (EOFError, KeyboardInterrupt):
                depth = "2"
            messages[0]["content"] = state["base_system"] + coach_addendum(depth)
            state["coach"] = True
            print(f"✅ 学习模式已开启（{COACH_DEPTHS.get(depth, COACH_DEPTHS['2'])[:12]}…）。说 /coach off 退出。")
    else:
        print(f"未知命令 {cmd}，输 /help 看全部命令。")
    return True


def main():
    ap = argparse.ArgumentParser(description="网文工坊 终端版")
    ap.add_argument("prompt", nargs="*", help="一次性输入；留空进入对话模式")
    ap.add_argument("--full", action="store_true", help="加载完整技能+知识库")
    ap.add_argument("--model", help="覆盖模型名")
    ap.add_argument("--list", action="store_true", help="列出所有技能")
    args = ap.parse_args()

    if args.list:
        list_skills()
        return

    try:
        from openai import OpenAI
    except ImportError:
        print("缺少依赖，请先运行： pip install openai")
        sys.exit(1)

    cfg = load_config()
    if args.model:
        cfg["model"] = args.model
    client = OpenAI(base_url=cfg["base_url"], api_key=cfg["api_key"])
    system = build_system_prompt(full=args.full)
    messages = [{"role": "system", "content": system}]

    print(f"网文工坊 终端版 | 模型 {cfg['model']} @ {cfg['base_url']}"
          + ("  [完整知识库]" if args.full else ""))

    if args.prompt:
        messages.append({"role": "user", "content": " ".join(args.prompt)})
        try:
            chat_stream(client, cfg["model"], messages)
        except Exception as e:
            print("\n[错误]", e)
        return

    state = {"base_system": system, "coach": False, "last_reply": ""}
    print_config(cfg)
    print()
    print(MODE_MENU)
    print("\n也可直接说人话开写。输 /help 看全部命令，exit 退出。\n")
    while True:
        try:
            user = input("你 > ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n再见。")
            break
        if user.lower() in ("exit", "quit", "退出"):
            break
        if not user:
            continue
        if user.startswith("/"):
            handle_command(user, cfg, client, messages, state)
            continue
        messages.append({"role": "user", "content": user})
        append_session_log(cfg, "你", user)
        print("工坊 > ", end="", flush=True)
        try:
            reply = chat_stream(client, cfg["model"], messages)
            messages.append({"role": "assistant", "content": reply})
            state["last_reply"] = reply
            append_session_log(cfg, "工坊", reply)
        except Exception as e:
            print("\n[错误]", e)
            messages.pop()


if __name__ == "__main__":
    main()
