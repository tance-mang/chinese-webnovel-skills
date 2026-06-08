# -*- coding: utf-8 -*-
"""
把 Markdown 草稿转成"供修改的 Word 草稿"。
用法：
  python build_docx.py 输入.md --title "书名" [--outdir 目录]

Markdown 约定：
  # 第一章 章名        -> 章标题（左对齐，进目录）
  正文段落...
  【修改：建议……】     -> 红色高亮的修改建议（行内或独立成段都可）

输出：《书名》修改N.docx（N 自动递增，绝不覆盖原文或已有"修改"文件），含目录。
"""
import sys, os, re, argparse

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SUGGEST_RE = re.compile(r'【修改[:：][^】]*】')
CHAPTER_RE = re.compile(r'^\s*#*\s*(第[0-9零一二三四五六七八九十百千两]+章.*)$')


def set_cjk(run, font='SimSun'):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:eastAsia'), font)


def add_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0, 0, 0)  # 纯黑，不要默认颜色
    set_cjk(r, 'SimHei')


def add_toc(doc):
    doc.add_paragraph().add_run('目录').bold = True
    p = doc.add_paragraph()
    run = p.add_run()
    for ftype, text in [('begin', None), (None, 'TOC \\o "1-3" \\h \\z \\u'), ('separate', None)]:
        pass
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    tnode = OxmlElement('w:t'); tnode.text = '（在 Word 里右键此处 → 更新域，即可生成页码目录）'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    for node in (fld_begin, instr, fld_sep, tnode, fld_end):
        run._r.append(node)
    doc.add_page_break()


def add_chapter(doc, text):
    p = doc.add_heading(level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 第X章 左对齐
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0, 0, 0)  # 章标题纯黑，不要 Word 默认的蓝色
    set_cjk(r, 'SimHei')


def add_body(doc, line):
    """正文段落；其中【修改：…】渲染成红色高亮。"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
    pos = 0
    for m in SUGGEST_RE.finditer(line):
        if m.start() > pos:
            r = p.add_run(line[pos:m.start()]); set_cjk(r)
        r = p.add_run(m.group(0))
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        r.bold = True
        set_cjk(r)
        pos = m.end()
    if pos < len(line):
        r = p.add_run(line[pos:]); set_cjk(r)


def pick_filename(title, outdir):
    """《书名》修改N.docx，N 从 1 起自增，绝不覆盖原文或已有修改稿。"""
    n = 1
    while True:
        name = f'《{title}》修改{n}.docx'
        path = os.path.join(outdir, name)
        if not os.path.exists(path):
            return path
        n += 1


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('input')
    ap.add_argument('--title', required=True)
    ap.add_argument('--outdir', default='.')
    args = ap.parse_args()

    with open(args.input, encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'SimSun'; style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    add_title(doc, args.title)
    add_toc(doc)

    body_chars = 0
    for line in lines:
        if not line.strip():
            continue
        cm = CHAPTER_RE.match(line)
        if cm:
            add_chapter(doc, cm.group(1).strip())
        else:
            add_body(doc, line)
            body_chars += len(SUGGEST_RE.sub('', line))

    os.makedirs(args.outdir, exist_ok=True)
    out = pick_filename(args.title, args.outdir)
    doc.save(out)

    # 字数提示（不含修改建议）
    tip = ''
    if body_chars < 8000:
        tip = '  ⚠ 低于短篇下限(8000)，可扩写'
    elif body_chars > 30000:
        tip = '  ⚠ 超过短篇上限(30000)，可拆分'
    elif not (10000 <= body_chars <= 20000):
        tip = '  (短篇正文建议草稿 10000–20000，留删/扩余地)'
    print(f'[OK] 已生成：{out}')
    print(f'     正文约 {body_chars} 字{tip}')
    print(f'     目录已插入（Word 里右键目录→更新域 生成页码）；【修改：…】为红色高亮建议')


if __name__ == '__main__':
    main()
